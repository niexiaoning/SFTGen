from typing import Any, Dict, List

from arborgraph.bases.base_reader import BaseReader

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "python-docx is required for reading Word documents. "
        "Please install it using: pip install python-docx"
    )


class DOCXReader(BaseReader):
    """
    Reader for Microsoft Word (.docx) files.
    Extracts text content from Word documents while preserving paragraph structure.
    Tables are converted to Markdown format for better structure preservation.
    """
    
    @staticmethod
    def _table_to_markdown(table) -> str:
        """
        Convert a Word table to Markdown format.
        
        :param table: docx table object
        :return: Markdown formatted table string
        """
        if len(table.rows) == 0:
            return ""
        
        # Extract all rows
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Get cell text, replace newlines with <br> or spaces, and escape pipe characters
                cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                cells.append(cell_text)
            # Only add non-empty rows
            if any(cell.strip() for cell in cells):
                rows.append(cells)
        
        if not rows:
            return ""
        
        # Determine number of columns (use max columns across all rows)
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols == 0:
            return ""
        
        # Normalize rows to have the same number of columns
        normalized_rows = []
        for row in rows:
            normalized_row = row + [""] * (max_cols - len(row))
            normalized_rows.append(normalized_row[:max_cols])  # Trim if too many columns
        
        # Build Markdown table
        markdown_lines = []
        
        # Header row (first row)
        if normalized_rows:
            header_row = normalized_rows[0]
            markdown_lines.append("| " + " | ".join(header_row) + " |")
            
            # Separator row
            markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
            
            # Data rows (rest of rows)
            for row in normalized_rows[1:]:
                markdown_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(markdown_lines)
    
    def _extract_text_from_element(self, element) -> str:
        """
        Recursively extract text from a docx element, including text in runs and nested elements.
        
        :param element: docx element (paragraph, table cell, etc.)
        :return: Extracted text content
        """
        text_parts = []
        
        # Extract text from runs (direct text content)
        if hasattr(element, 'runs'):
            for run in element.runs:
                if run.text:
                    text_parts.append(run.text)
        
        # Extract text from paragraphs if element has them
        if hasattr(element, 'paragraphs'):
            for para in element.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text_parts.append(para_text)
        
        # Extract text from tables if element has them
        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = self._extract_text_from_element(cell)
                        if cell_text.strip():
                            text_parts.append(cell_text)
        
        return "\n".join(text_parts)
    
    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Read text content from a .docx file.
        Extracts text from main document paragraphs and tables only.
        Excludes headers, footers, footnotes, endnotes, and page numbers.
        
        :param file_path: Path to the .docx file.
        :return: List containing a single dictionary with the document content.
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from main document paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    paragraphs.append(para_text)
            
            if paragraphs:
                text_parts.append("\n\n".join(paragraphs))
            
            # Extract text from tables in main document
            tables_text = []
            for table in doc.tables:
                if len(table.rows) == 0:
                    continue
                
                try:
                    # Try to convert table to Markdown format
                    markdown_table = self._table_to_markdown(table)
                    if markdown_table:
                        tables_text.append(markdown_table)
                except Exception:
                    # If Markdown conversion fails, fallback to plain text format
                    try:
                        table_rows = []
                        for row in table.rows:
                            row_cells = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_cells.append(cell_text)
                            if row_cells:
                                table_rows.append(" | ".join(row_cells))
                        if table_rows:
                            tables_text.append("\n".join(table_rows))
                    except Exception:
                        # If even plain text extraction fails, try recursive extraction
                        try:
                            table_text = self._extract_text_from_element(table)
                            if table_text.strip():
                                tables_text.append(table_text)
                        except Exception:
                            pass
            
            if tables_text:
                text_parts.append("\n\n".join(tables_text))
            
            # If no text found in standard locations, try comprehensive extraction from document body only
            if not text_parts or not any(part.strip() for part in text_parts):
                # Try to extract all text from document body XML (excluding headers/footers)
                try:
                    from docx.oxml.text.paragraph import CT_P
                    from docx.oxml.table import CT_Tbl
                    from docx.text.paragraph import Paragraph
                    from docx.table import Table
                    
                    body_elements = []
                    # Iterate through all elements in document body only (not headers/footers)
                    for element in doc.element.body:
                        if isinstance(element, CT_P):
                            # Extract text from paragraph
                            try:
                                para = Paragraph(element, doc)
                                para_text = para.text.strip()
                                if para_text:
                                    body_elements.append(para_text)
                            except Exception:
                                # If Paragraph creation fails, try direct XML text extraction
                                try:
                                    para_text = element.text.strip()
                                    if para_text:
                                        body_elements.append(para_text)
                                except Exception:
                                    pass
                        elif isinstance(element, CT_Tbl):
                            # Extract text from table
                            try:
                                table = Table(element, doc)
                                # Try to extract text from all cells
                                table_texts = []
                                for row in table.rows:
                                    row_texts = []
                                    for cell in row.cells:
                                        cell_text = cell.text.strip()
                                        if cell_text:
                                            row_texts.append(cell_text)
                                    if row_texts:
                                        table_texts.append(" | ".join(row_texts))
                                if table_texts:
                                    body_elements.append("\n".join(table_texts))
                            except Exception:
                                # If Table creation fails, try recursive text extraction from element
                                try:
                                    table_text = self._extract_text_from_element(element)
                                    if table_text.strip():
                                        body_elements.append(table_text)
                                except Exception:
                                    pass
                    
                    if body_elements:
                        text_parts = ["\n\n".join(body_elements)]
                except Exception:
                    # Last resort: try to extract all text using XML text extraction from body only
                    try:
                        all_text = []
                        # Extract all text nodes from document body only
                        for para in doc.element.body.iter():
                            if para.text and para.text.strip():
                                all_text.append(para.text.strip())
                        if all_text:
                            text_parts = ["\n\n".join(all_text)]
                    except Exception:
                        pass
            
            # Combine all text parts
            full_text = "\n\n".join(text_parts) if text_parts else ""
            
            docs = [{
                self.text_column: full_text,
                "type": "text"
            }]
            
            return self.filter(docs)
        except Exception as e:
            raise ValueError(f"Failed to read Word document {file_path}: {str(e)}")

